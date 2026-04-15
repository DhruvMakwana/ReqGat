"""
Document generation service — produces BRD and FRD in .docx and PDF formats.
"""
from __future__ import annotations

import os
import uuid
from datetime import datetime, timezone
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.config import settings
from app.models.project import Project
from app.models.requirement import Requirement
from app.models.scenario import Scenario


def _ensure_storage(project_id: uuid.UUID) -> Path:
    path = Path(settings.STORAGE_PATH) / str(project_id)
    path.mkdir(parents=True, exist_ok=True)
    return path


def _add_heading(doc: DocxDocument, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_paragraph(doc: DocxDocument, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def generate_brd(
    project: Project,
    requirements: list[Requirement],
) -> Path:
    """Generate a Business Requirement Document (.docx)."""
    doc = DocxDocument()

    # Title
    title = doc.add_heading(f"Business Requirement Document", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Project: {project.name}")
    doc.add_paragraph(f"Domain: {project.domain_type.upper()}")
    doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph("")

    # 1. Project Overview
    _add_heading(doc, "1. Project Overview", 1)
    doc.add_paragraph(project.description or "No description provided.")

    # 2. Objectives
    _add_heading(doc, "2. Objectives", 1)
    doc.add_paragraph(
        f"To implement a {project.domain_type.upper()} solution that meets all defined "
        "business requirements with clear scope boundaries and scenario coverage."
    )

    # 3. Scope
    _add_heading(doc, "3. Scope", 1)

    in_scope = [r for r in requirements if r.category == "what_to_do"]
    out_scope = [r for r in requirements if r.category == "what_not_to_do"]

    _add_heading(doc, "3.1 In Scope", 2)
    if in_scope:
        for req in in_scope:
            doc.add_paragraph(f"• [{req.unique_id}] {req.title}", style="List Bullet")
    else:
        doc.add_paragraph("No in-scope items defined.")

    _add_heading(doc, "3.2 Out of Scope", 2)
    if out_scope:
        for req in out_scope:
            doc.add_paragraph(f"• [{req.unique_id}] {req.title}", style="List Bullet")
    else:
        doc.add_paragraph("No exclusions defined.")

    # 4. High-Level Requirements
    _add_heading(doc, "4. High-Level Requirements", 1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "ID"
    hdr[1].text = "Requirement"
    hdr[2].text = "Category"
    hdr[3].text = "Priority"

    for req in requirements:
        row = table.add_row().cells
        row[0].text = req.unique_id
        row[1].text = req.title
        row[2].text = req.category.replace("_", " ").title()
        row[3].text = req.priority.title()

    storage = _ensure_storage(project.id)
    file_path = storage / "brd.docx"
    doc.save(str(file_path))
    return file_path


def generate_frd(
    project: Project,
    requirements: list[Requirement],
) -> Path:
    """Generate a Functional Requirement Document (.docx)."""
    doc = DocxDocument()

    title = doc.add_heading("Functional Requirement Document", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Project: {project.name}")
    doc.add_paragraph(f"Domain: {project.domain_type.upper()}")
    doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph("")

    _add_heading(doc, "1. Detailed Requirements", 1)

    for req in requirements:
        _add_heading(doc, f"[{req.unique_id}] {req.title}", 2)
        doc.add_paragraph(f"Category: {req.category.replace('_', ' ').title()}")
        doc.add_paragraph(f"Priority: {req.priority.title()}")
        doc.add_paragraph(f"Status: {req.status.title()}")
        if req.description:
            doc.add_paragraph(f"Description: {req.description}")

        accepted_scenarios = [s for s in req.scenarios if s.status == "accepted"]
        if accepted_scenarios:
            _add_heading(doc, "Scenarios", 3)
            for sc in accepted_scenarios:
                doc.add_paragraph(
                    f"• [{sc.type.replace('_', ' ').title()}] {sc.description}",
                    style="List Bullet",
                )
        doc.add_paragraph("")

    # Conditional logic summary
    _add_heading(doc, "2. What-If Conditions", 1)
    what_if_reqs = [r for r in requirements if r.category == "what_if"]
    if what_if_reqs:
        for req in what_if_reqs:
            doc.add_paragraph(f"• {req.title}: {req.description or ''}", style="List Bullet")
    else:
        doc.add_paragraph("No what-if conditions defined.")

    # Exception handling
    _add_heading(doc, "3. Exception Handling", 1)
    exception_scenarios = [
        (req, sc)
        for req in requirements
        for sc in req.scenarios
        if sc.type == "exception" and sc.status == "accepted"
    ]
    if exception_scenarios:
        for req, sc in exception_scenarios:
            doc.add_paragraph(
                f"• [{req.unique_id}] {sc.description}", style="List Bullet"
            )
    else:
        doc.add_paragraph("No exception scenarios defined.")

    storage = _ensure_storage(project.id)
    file_path = storage / "frd.docx"
    doc.save(str(file_path))
    return file_path


def generate_pdf_from_docx(docx_path: Path) -> Path:
    """Convert a .docx to PDF using WeasyPrint via HTML intermediate."""
    # Build a simple HTML representation of the document content
    from docx import Document as DocxDocument
    doc = DocxDocument(str(docx_path))

    html_parts = ["<html><head><meta charset='utf-8'><style>"]
    html_parts.append("""
        body { font-family: Arial, sans-serif; margin: 40px; font-size: 12pt; }
        h1 { color: #1a365d; border-bottom: 2px solid #1a365d; }
        h2 { color: #2b6cb0; }
        h3 { color: #2c5282; }
        table { border-collapse: collapse; width: 100%; margin: 10px 0; }
        td, th { border: 1px solid #cbd5e0; padding: 6px 10px; }
        th { background: #2b6cb0; color: white; }
        p { margin: 4px 0; }
    """)
    html_parts.append("</style></head><body>")

    for para in doc.paragraphs:
        style = para.style.name
        text = para.text.replace("<", "&lt;").replace(">", "&gt;")
        if not text.strip():
            html_parts.append("<br>")
            continue
        if "Heading 1" in style:
            html_parts.append(f"<h1>{text}</h1>")
        elif "Heading 2" in style:
            html_parts.append(f"<h2>{text}</h2>")
        elif "Heading 3" in style:
            html_parts.append(f"<h3>{text}</h3>")
        elif "List Bullet" in style:
            html_parts.append(f"<p>{text}</p>")
        else:
            html_parts.append(f"<p>{text}</p>")

    for table in doc.tables:
        html_parts.append("<table>")
        for i, row in enumerate(table.rows):
            html_parts.append("<tr>")
            for cell in row.cells:
                tag = "th" if i == 0 else "td"
                html_parts.append(f"<{tag}>{cell.text}</{tag}>")
            html_parts.append("</tr>")
        html_parts.append("</table>")

    html_parts.append("</body></html>")
    html_content = "".join(html_parts)

    pdf_path = docx_path.with_suffix(".pdf")
    try:
        import weasyprint
        weasyprint.HTML(string=html_content).write_pdf(str(pdf_path))
    except Exception:
        # Fallback: write HTML if WeasyPrint fails (missing system deps)
        html_path = docx_path.with_suffix(".html")
        html_path.write_text(html_content, encoding="utf-8")
        pdf_path = html_path  # return HTML path as fallback
    return pdf_path
